#!/usr/bin/env python3
"""Generate SWES Private Markets Stress Lens FRD"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Page setup
for s in doc.sections:
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

# Style setup
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for lvl, sz in [(1, 22), (2, 16), (3, 13), (4, 11)]:
    h = doc.styles[f'Heading {lvl}']
    h.font.name = 'Calibri'
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C) if lvl <= 2 else RGBColor(0x2E, 0x75, 0xB6)

def shade_cell(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.name = 'Calibri'
        shade_cell(c, '1B3A5C')
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(str(v))
            r.font.size = Pt(9); r.font.name = 'Calibri'
            if ri % 2 == 1: shade_cell(c, 'F2F7FB')
    return t

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F5F5F5'); shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(8.5)

def eq(text, num=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Cambria Math'; r.font.size = Pt(11); r.italic = True
    if num:
        r2 = p.add_run(f"    ({num})")
        r2.font.size = Pt(10); r2.bold = True

# ═══════ COVER PAGE ═══════
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FUNCTIONAL REQUIREMENTS DOCUMENT'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x2E,0x75,0xB6); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SWES Private Markets Stress Lens'); r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1B,0x3A,0x5C); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Multi-Agent System-Wide Liquidity & Credit Stress Simulator'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Proof of Concept \u2014 Build Specification  |  Version 1.0'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'{datetime.date.today().strftime("%B %d, %Y")}  |  CONFIDENTIAL'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x99,0x99,0x99)

doc.add_page_break()

# ═══════ 1. EXECUTIVE SUMMARY ═══════
doc.add_heading('1. Executive Summary & Objectives', level=1)
doc.add_paragraph('This document specifies the Functional Requirements for the "SWES Private Markets Stress Lens" (the "Tool"), a proof-of-concept interactive stress testing application for the Bank of England\'s second SWES exercise focused on private markets. The Tool adapts the Van den End (2009) DNB Liquidity Stress-Tester framework from a homogeneous banking system to a heterogeneous multi-agent private markets ecosystem.')

doc.add_heading('Product Vision', level=3)
doc.add_paragraph('An interactive Streamlit application demonstrating how a severe macroeconomic downturn propagates through the private markets ecosystem \u2014 from direct portfolio impacts to system-wide behavioural feedback effects \u2014 showing that rational individual actions by banks, AAMs, and institutional investors collectively amplify the initial shock.')

doc.add_heading('Key Objectives', level=3)
doc.add_paragraph('1. Demonstrate understanding of the SWES exercise\'s system-wide behavioural focus, distinguishing it from conventional single-entity stress testing.\n2. Show a working multi-agent stress model where agents (bank, AAM, pension fund) respond to a common shock and their responses feed back on each other.\n3. Quantify the "amplification ratio" \u2014 the gap between naive direct-stress losses and system-aware losses.\n4. Use real public data for portfolio construction; calibrate behavioural parameters to published SWES 1 findings.\n5. Create a reusable analytical engine that compounds in value with each client engagement.')

doc.add_heading('Timeline', level=3)
add_table(['Milestone', 'Target', 'Deliverable'], [
    ['Internal Demo (v0.1)', 'Week 2', 'Core engine + Streamlit UI, synthetic data'],
    ['Data Integration', 'Week 3', 'Real public data, SWES 1 calibration'],
    ['Client Demo (v1.0)', 'Week 4', 'Polished UI, dual narrative (bank + AAM)'],
    ['Enhancement', 'Week 5-6', 'VM simulator, agent deepening, React eval'],
])

doc.add_page_break()

# ═══════ 2. THEORETICAL FOUNDATION ═══════
doc.add_heading('2. Theoretical Foundation', level=1)

doc.add_heading('2.1 SWES Exercise Context', level=2)
doc.add_paragraph('The Bank of England\'s SWES exercise asks: "When multiple firms act rationally in their own self-interest during a stress, do their collective actions amplify the shock?" The first SWES (2023-2024) demonstrated they do. The second SWES focuses on private markets: PE, private credit, leveraged loans, CLOs, and high-yield bonds. Scenario phase launching early 2026, final report early 2027.')

doc.add_heading('SWES 1 Calibration Anchors', level=3)
add_table(['Finding', 'Quantitative Detail', 'Model Relevance'], [
    ['Aggregate NBFI margin calls', '~\u00a394 billion', 'Liquidity drain magnitude'],
    ['LDI recapitalisation requests', '\u00a316.5 billion to pension schemes', 'Cross-agent capital call dynamics'],
    ['NBFI gilt sales under stress', '~\u00a34.7 billion', 'Selling pressure calibration'],
    ['Bank capacity consumed', '70%; +\u00a30.5bn long gilts exhausts it', 'Bank tightening threshold'],
    ['Repo financing mismatch', '>1/3 of NBFIs refused by all banks', 'Financing withdrawal severity'],
    ['IM estimation error', 'Systematic overestimation by banks/NBFIs', 'Margin call uncertainty'],
])

doc.add_heading('SWES 2 Sizing Data', level=3)
doc.add_paragraph('Global PE/PC AUM: ~$11 trillion (from $3T a decade ago). UK PE-sponsored corporates: up to 15% of UK corporate debt, 10% of UK private sector employment (>2 million jobs). Over 80% of PE-sponsored corporate debt is in alternative credit lines vs. 30% for UK corporates broadly.')

doc.add_heading('2.2 Van den End Framework', level=2)
doc.add_paragraph('Van den End (2009) developed a liquidity stress-testing model at DNB with first-round and second-round (feedback) effects, including behavioural reactions and reputation effects. Applied to 82 Dutch banks. Three-stage architecture:')

add_table(['Stage', 'Description', 'Key Equation', 'Output'], [
    ['1: First Round', 'Direct scenario impact via stressed haircuts and run-off rates', 'B\u2081 = B\u2080 \u2212 E\u2081, where E\u2081 = \u03a3 I\u1d62 \u00d7 w\u2081,\u1d62', 'Buffer after shock (B\u2081)'],
    ['2: Reactions', 'Agents with buffer decline > threshold \u03b8 take mitigating actions', 'RI\u1d62 = (B\u2080 \u2212 B\u2081) \u00d7 (I\u1d62 / \u03a3I\u2c7c)', 'Buffer after reactions (B\u2082)'],
    ['3: Feedback', 'Collective reactions + reputation risk create second-round stress', 'w\u2082 = f(w\u2081, \u03a3q, similarity, s); w*\u2082 = w\u2082 \u00d7 \u221as', 'Final buffer (B\u2083)'],
])

doc.add_paragraph('Key finding: In the credit crisis scenario, first round erased 13% of buffer. After second round: 39% total. Amplification ratio ~3.0x. Second round effects were 2x the first round. More than half of second round impact on reacting banks was reputation risk.')

doc.add_heading('2.3 Our Adaptation', level=2)
doc.add_paragraph('We extend Van den End in three ways:\n(a) Heterogeneous agent types: Bank, AAM, Pension Fund (not homogeneous banks).\n(b) Cross-agent contagion channels: bank\u2192AAM (lending withdrawal), AAM\u2192PF (valuation markdowns), PF\u2192AAM (commitment reductions), all\u2192credit markets.\n(c) Prolonged downturn scenario (not 10-day market shock): credit deterioration \u2192 default cascades \u2192 financing withdrawal.')

doc.add_page_break()

# ═══════ 3. SYSTEM ARCHITECTURE ═══════
doc.add_heading('3. System Architecture', level=1)

doc.add_heading('3.1 Module Decomposition', level=2)
code("""
Data Flow: data_loader \u2192 agents \u2192 scenario_engine \u2192 feedback_loop \u2192 visualisation \u2192 Streamlit UI

Modules:
\u2502 data_loader.py      \u2502 Load portfolio CSV, scenario JSON, calibration anchors
\u2502 agents.py            \u2502 Agent archetypes, balance sheets, thresholds, reactions
\u2502 scenario_engine.py   \u2502 Macro scenario \u2192 company stress \u2192 Stage 1 impacts
\u2502 feedback_loop.py     \u2502 Stage 2 reactions, Stage 3 feedback, reputation risk
\u2502 visualisation.py     \u2502 Plotly charts and dashboard components
\u2502 app.py               \u2502 Streamlit entry point, UI layout, demo flow
""")

doc.add_page_break()

# ═══════ 4. DATA MODEL ═══════
doc.add_heading('4. Data Model & Structures', level=1)

doc.add_heading('4.1 Agent Archetypes & Balance Sheets', level=2)
doc.add_paragraph('Each agent has a stylized balance sheet with items relevant to private markets stress. Proportions calibrated from BoE aggregate statistics and SWES 2 sizing data.')

doc.add_heading('Agent Type 1: Bank', level=3)
add_table(['Balance Sheet Item', 'Category', 'Stress Sensitivity', 'Notes'], [
    ['Leveraged Loan Book', 'Asset', 'EBITDA decline \u2192 ICR stress \u2192 PD', 'Core exposure; real company data'],
    ['Subscription Credit Lines', 'Asset (committed)', 'LP drawdown risk', 'Drawn by PE funds for bridging'],
    ['NAV Lending Facilities', 'Asset', 'Collateral (fund NAV) markdowns', 'Secured against fund NAV'],
    ['CLO Tranches', 'Asset', 'Spread widening; OC test breaches', 'AAA through equity tranches'],
    ['Revolving Credit Facilities', 'Asset', 'Drawdown by stressed corporates', 'Committed but undrawn'],
    ['Wholesale Funding', 'Liability', 'Run-off rate under stress', 'Interbank, CP, bonds'],
    ['Capital Buffer (CET1)', 'Equity', 'Absorbs credit losses', 'Starting CET1 ~14%'],
])

doc.add_heading('Agent Type 2: AAM', level=3)
add_table(['Balance Sheet Item', 'Category', 'Stress Sensitivity', 'Notes'], [
    ['PE Portfolio (UK companies)', 'Asset', 'EBITDA \u00d7 multiple compression', 'Real UK PE-backed companies'],
    ['Private Credit Fund', 'Asset', 'Default rates, recovery rates', 'Loans to UK corporates'],
    ['Fund-Level Leverage', 'Liability', 'Interest coverage, covenants', 'Sub lines, NAV lines drawn'],
    ['Dry Powder', 'Asset (potential)', 'LP willingness to fund calls', 'Rescue capital capacity'],
    ['LP Commitments', 'Liability (future)', 'LP default/reduction risk', 'Vintage-year sensitivity'],
])

doc.add_heading('Agent Type 3: Pension Fund', level=3)
add_table(['Balance Sheet Item', 'Category', 'Stress Sensitivity', 'Notes'], [
    ['Liquid Portfolio (Gilts)', 'Asset', 'Mark-to-market on rate shock', 'Core liquid buffer'],
    ['Corporate Bond Holdings', 'Asset', 'Spread widening; illiquidity', 'Sterling IG'],
    ['PE/PC Allocation', 'Asset (illiquid)', 'NAV markdown', 'Cannot be liquidated quickly'],
    ['LDI Hedge Positions', 'Asset/Liability', 'Margin calls on rate shock', 'Key liquidity drain channel'],
    ['Capital Call Obligations', 'Liability (contingent)', 'Forced selling to fund calls', 'Arrives at worst time'],
])

doc.add_heading('4.2 Scenario Parameters', level=2)
add_table(['Parameter', 'Symbol', 'Default', 'Slider Range', 'Calibration'], [
    ['GDP Decline', '\u0394_GDP', '-3.0%', '0% to -8%', 'GFC = -6%; BoE stress scenarios'],
    ['Base Rate Change', '\u0394_rate', '+150 bps', '-100 to +300', 'SWES 1: +115 bps gilts'],
    ['IG Spread Widening', '\u0394_cs_ig', '+130 bps', '0 to +300', 'SWES 1: +130 bps sterling IG'],
    ['Lev Loan Spread', '\u0394_cs_lev', '+400 bps', '0 to +800', 'GFC ~+600; COVID ~+350'],
    ['Equity Decline', '\u0394_eq', '-25%', '0% to -50%', 'SWES 1: < historical max'],
    ['PE Multiple Compression', '\u0394_mult', '-2.0x', '0 to -4.0x', 'GFC: ~2-3x from peaks'],
    ['Market Stress (s)', 's', '2.0', '1.0 to 3.0', 'Van den End: VIX-calibrated'],
])

doc.add_heading('4.3 Contagion Channel Definitions', level=2)
add_table(['Channel', 'From \u2192 To', 'Mechanism', 'SWES Evidence'], [
    ['C1', 'Scenario \u2192 All', 'Direct macro shock', 'SWES 1 scenario design'],
    ['C2', 'Bank \u2192 AAM', 'Lending tightening, refuse new sub lines', '>1/3 NBFIs refused repo'],
    ['C3', 'AAM \u2192 Portfolio Cos', 'Reduced support, forced restructuring', 'SWES 2: real economy focus'],
    ['C4', 'Pension \u2192 AAM', 'LP reduces commitments, slows capital calls', 'Pension fund stress dynamics'],
    ['C5', 'AAM \u2192 Credit Markets', 'Forced CLO/lev loan selling', 'Corp bond "jump to illiquidity"'],
    ['C6', 'Credit Markets \u2192 Bank', 'Wider spreads, MTM losses', 'Banks limited risk appetite'],
    ['C7', 'Reputation \u2192 Reacting Agent', 'Signalling effect (\u221as)', '>50% of 2nd round effect'],
])

doc.add_page_break()

# ═══════ 5. MATHEMATICAL MODEL ═══════
doc.add_heading('5. Mathematical Model \u2014 Core Equations', level=1)
doc.add_paragraph('Adapted from Van den End (2009). Superscript a = agent type, subscript i = balance sheet item. Deterministic weights (not Monte Carlo) for POC.')

doc.add_heading('5.1 Stage 1: First Round Effects', level=2)
eq('B\u2080\u1d43 = \u03a3\u1d62 I_non_cal,i\u1d43', '1')
doc.add_paragraph('Initial buffer: CET1 for bank, unrealised NAV + dry powder for AAM, liquid assets net of near-term liabilities for pension fund.')
eq('E\u2081\u1d43 = \u03a3\u1d62 I\u1d62\u1d43 \u00d7 w\u2081,\u1d62\u1d43(scenario_params)', '2')
eq('B\u2081\u1d43 = B\u2080\u1d43 \u2212 E\u2081\u1d43', '3')

doc.add_heading('Company-Level Stress Chain', level=3)
eq('EBITDA_stressed = EBITDA_base \u00d7 (1 + \u03b2_sector \u00d7 \u0394_GDP)', '2b')
eq('Interest_stressed = Debt \u00d7 (base_rate + spread + \u0394_rate + \u0394_cs_lev)', '2d')
eq('ICR_stressed = EBITDA_stressed / Interest_stressed', '2c')
eq('PD_stressed = \u03a6((ICR_trigger \u2212 ICR_stressed) / \u03c3_ICR)', '2e')
eq('Valuation_stressed = EBITDA_stressed \u00d7 (Multiple_base + \u0394_mult)', '2f')

doc.add_heading('Sector GDP Elasticities (\u03b2)', level=3)
add_table(['Sector', '\u03b2 (% EBITDA decline per 1% GDP decline)', 'Source'], [
    ['Technology', '1.5', 'ONS GVA data, GFC/COVID calibration'],
    ['Healthcare', '0.8', 'Defensive sector'],
    ['Industrials', '2.5', 'Highly cyclical'],
    ['Consumer Discretionary', '3.0', 'Most cyclical'],
    ['Consumer Staples', '0.5', 'Defensive'],
    ['Business Services', '1.8', 'Moderately cyclical'],
    ['Financial Services', '2.0', 'Pro-cyclical'],
    ['Energy', '2.2', 'Commodity-sensitive'],
    ['Telecom', '1.0', 'Stable'],
    ['Real Estate', '2.8', 'Rate and cycle sensitive'],
])

doc.add_heading('5.2 Stage 2: Agent Reactions', level=2)
eq('q\u1d43 = 1 if (E\u2081\u1d43 / B\u2080\u1d43) > \u03b8\u1d43, else 0', '4')

add_table(['Agent Type', '\u03b8', 'Rationale'], [
    ['Bank', '0.40', 'Van den End empirical (Dutch banks); regulatory buffer thresholds'],
    ['AAM', '0.25', 'Less experienced with stress; react earlier; GP reputation incentives'],
    ['Pension Fund', '0.50', 'Slower governance; fiduciary inertia; post-2022 LDI buffers improved'],
])

eq('RI\u1d62\u1d43 = (B\u2080\u1d43 \u2212 B\u2081\u1d43) \u00d7 (I\u1d62\u1d43 / \u03a3\u2c7c I\u2c7c\u1d43)', '5')
eq('B\u2082\u1d43 = B\u2081\u1d43 + \u03a3\u1d62 RI\u1d62\u1d43 \u00d7 (1 \u2212 w\u2081,\u1d62\u1d43 / 100)', '6')

doc.add_heading('5.3 Stage 3: Second Round / Feedback', level=2)
doc.add_paragraph('Decomposed into explicit cross-agent channel effects:')
eq('E\u2082,bank_tightening = intensity_C2 \u00d7 s \u00d7 (total_bank_reactions / total_PM_exposure)', '7a')
eq('E\u2082,market_freeze = intensity_C5 \u00d7 s \u00d7 (total_credit_sales / market_depth)', '7b')
eq('E\u2082,reputation = (w\u2082 \u2212 w\u2081) \u00d7 \u221as    [reacting agents only, Van den End eq 7]', '7c')
eq('B\u2083\u1d43 = B\u2082\u1d43 \u2212 E\u2082\u1d43', '8')

doc.add_heading('5.4 Amplification Ratio', level=2)
eq('Amplification Ratio = (B\u2080 \u2212 B\u2083) / (B\u2080 \u2212 B\u2081)', '9')
doc.add_paragraph('Target range: 1.5\u20133.0x. Van den End credit crisis = ~3.0x. This is the central "aha moment" metric.')

doc.add_page_break()

# ═══════ 6. FUNCTION SIGNATURES ═══════
doc.add_heading('6. Function Signatures & Module Specifications', level=1)

doc.add_heading('6.1 data_loader.py', level=2)
code("""@dataclass
class Company:
    name: str
    sector: str                    # SIC-based classification
    pe_sponsor: str
    revenue_mm: float
    ebitda_mm: float
    total_debt_mm: float
    interest_expense_mm: float
    employee_count: int
    debt_ebitda: float             # Derived
    icr: float                     # Derived
    maturity_years: float
    is_floating_rate: bool         # Default True for LBO debt
    covenant_icr_trigger: float    # Default 1.25

@dataclass
class ScenarioParams:
    gdp_decline_pct: float         # e.g., -3.0
    base_rate_change_bps: float    # e.g., 150
    ig_spread_widen_bps: float     # e.g., 130
    lev_loan_spread_widen_bps: float  # e.g., 400
    equity_decline_pct: float      # e.g., -25.0
    pe_multiple_compression: float # e.g., -2.0
    market_stress_s: float         # 1.0-3.0

SECTOR_ELASTICITIES = {
    'Technology': 1.5, 'Healthcare': 0.8, 'Industrials': 2.5,
    'Consumer_Discretionary': 3.0, 'Consumer_Staples': 0.5,
    'Business_Services': 1.8, 'Financial_Services': 2.0,
    'Energy': 2.2, 'Telecom': 1.0, 'Real_Estate': 2.8,
}

def load_companies(csv_path: str) -> List[Company]: ...
def load_default_scenario() -> ScenarioParams: ...""")

doc.add_heading('6.2 agents.py', level=2)
code("""class AgentType(Enum):
    BANK = "bank"
    AAM = "aam"
    PENSION_FUND = "pension_fund"

@dataclass
class BalanceSheetItem:
    name: str
    category: str        # 'asset', 'liability', 'equity', 'contingent'
    amount_mm: float
    w1: float = 0.0      # First-round stress weight (0-100)
    w2: float = 0.0      # Second-round stress weight (0-100)
    is_reaction_instrument: bool = False

@dataclass
class Agent:
    name: str
    agent_type: AgentType
    balance_sheet: List[BalanceSheetItem]
    theta: float         # Reaction threshold
    B0, B1, B2, B3: float = 0.0
    E1, E2: float = 0.0
    has_reacted: bool = False
    reactions: Dict[str, float] = field(default_factory=dict)

    def compute_buffer(self) -> float: ...
    def should_react(self) -> bool:
        return (self.E1 / self.B0) > self.theta if self.B0 > 0 else False
    def compute_reactions(self) -> Dict[str, float]: ...

# Factory functions:
def create_bank_agent(total_pm_exposure_mm=5000.0, ...) -> Agent: ...
def create_aam_agent(total_aum_mm=15000.0, ...) -> Agent: ...
def create_pension_fund_agent(total_aum_mm=10000.0, ...) -> Agent: ...""")

doc.add_heading('6.3 scenario_engine.py', level=2)
code("""@dataclass
class CompanyStressResult:
    company: Company
    ebitda_stressed_mm: float
    icr_stressed: float
    pd_stressed: float
    covenant_breached: bool
    valuation_base_mm: float
    valuation_stressed_mm: float
    valuation_markdown_pct: float
    employees_at_risk: int

def stress_company(company: Company, scenario: ScenarioParams,
                   base_ev_ebitda_multiple: float = 8.0) -> CompanyStressResult:
    # Implements equations 2b-2f
    sector_beta = SECTOR_ELASTICITIES.get(company.sector, 2.0)
    ebitda_stressed = company.ebitda_mm * (1 + sector_beta * scenario.gdp_decline_pct/100)
    # ... interest stress, ICR, PD, valuation ...

def apply_scenario_to_agent(agent: Agent, scenario: ScenarioParams,
                            companies: List[Company] = None) -> Agent: ...
def compute_real_economy_impact(results: List[CompanyStressResult]) -> Dict: ...""")

doc.add_heading('6.4 feedback_loop.py', level=2)
code("""@dataclass
class ContagionChannel:
    channel_id: str
    from_agent_type: AgentType
    to_agent_type: AgentType
    mechanism: str
    intensity: float         # 0.0 to 1.0
    affects_items: List[str]

def compute_second_round_weights(agents, channels, s) -> Dict: ...
def apply_reputation_risk(agent, s) -> float:
    # Van den End eq 7: w*2 = w2 x sqrt(s)
    return sum(agent.reactions.values()) * (sqrt(s) - 1.0) if agent.has_reacted else 0.0

def compute_amplification_ratio(agents) -> Dict[str, float]:
    # Eq 9: (B0-B3)/(B0-B1) per agent + system-wide weighted avg

def run_full_simulation(agents, scenario, channels=None, n_iterations=3) -> Dict:
    # Master: Stage1 -> reactions -> Stage2 -> feedback -> Stage3 -> amplification""")

doc.add_heading('6.5 visualisation.py', level=2)
code("""def plot_agent_buffers_waterfall(agents) -> go.Figure:
    # B0 -> B1 -> B2 -> B3 waterfall; gap between B1 and B3 = amplification
    # THIS IS THE CORE 'AHA MOMENT' CHART

def plot_portfolio_heatmap(company_results) -> go.Figure:
    # Companies x metrics (ICR, PD, covenant breach) heatmap by sector

def plot_contagion_sankey(agents, channels) -> go.Figure:
    # Sankey: stress flows between agents, width = magnitude

def plot_amplification_gauge(ratios) -> go.Figure:
    # Gauge: Green 1-1.5, Yellow 1.5-2.0, Orange 2.0-2.5, Red 2.5+

def plot_real_economy_impact(impact) -> go.Figure:
    # Employees at risk by sector, credit supply reduction

def plot_scenario_comparison(direct, system) -> go.Figure:
    # Side-by-side: Act 2 vs Act 3 comparison""")

doc.add_page_break()

# ═══════ 7. STREAMLIT UI ═══════
doc.add_heading('7. Streamlit UI Specification', level=1)

doc.add_heading('7.1 Layout', level=2)
code("""# app.py structure
# SIDEBAR: Scenario sliders + "Enable System-Wide Feedback" toggle
# MAIN: 4 tabs
#   Tab 1: Portfolio & Agents (balance sheet overview)
#   Tab 2: Direct Stress (company heatmap, agent buffers B0->B1)
#   Tab 3: System-Wide Dynamics (contagion Sankey, amplification gauge, B0->B3)
#   Tab 4: Real Economy Impact (employment, credit supply, BoE reporting)""")

doc.add_heading('7.2 Demo Flow (4-Act Structure)', level=2)
add_table(['Act', 'Time', 'Tab', 'What Happens', 'Audience Emotion'], [
    ['1: Setup', '5m', '1', 'Show portfolio of real UK PE companies. 3 agent balance sheets.', 'Comfort: "they know our world"'],
    ['2: Direct', '7m', '2', 'Adjust sliders. EBITDA \u2192 ICR \u2192 covenants \u2192 defaults. Feedback OFF.', 'Mild concern: "standard stress test"'],
    ['3: System', '8m', '3', 'FLIP TOGGLE. Losses double. Contagion Sankey. Amplification 2-2.5x.', 'Shock: "losses just doubled"'],
    ['4: Real Econ', '5m', '4', 'Employment at risk. Credit supply gap. "What BoE wants in your submission."', 'Urgency: "we need help"'],
])

doc.add_page_break()

# ═══════ 8. DATA SOURCING ═══════
doc.add_heading('8. Data Sourcing Plan', level=1)

doc.add_heading('WS1: UK PE Company Universe (Analyst 1, Days 1-3)', level=3)
doc.add_paragraph('Goal: 40-50 real UK PE-backed companies with financials.\n\nStep 1: Identify companies from SWES 2 participant AAM portfolio pages (Blackstone, KKR, Carlyle, Hg, Permira list UK companies on websites), BVCA annual reports, news.\nStep 2: Pull financials from Companies House: revenue, operating profit, total debt, interest payable, employees.\nStep 3: Classify by sector (SIC codes \u2192 our taxonomy).\nStep 4: Add PE sponsor.\n\nOutput CSV: company_name, pe_sponsor, sector, revenue_mm, ebitda_mm, total_debt_mm, interest_expense_mm, employee_count, filing_year, is_floating_rate.\n\nFallback: Synthetic companies calibrated from Preqin/BVCA sector averages.')

doc.add_heading('WS2: Scenario Calibration (Analyst 2, Days 1-3)', level=3)
doc.add_paragraph('Goal: Sector-specific EBITDA-GDP elasticities.\n\nStep 1: ONS sectoral GVA data 2006-2023 (GFC + COVID).\nStep 2: Compute \u03b2 per sector. Cross-ref with S&P/Moody\'s default studies.\nStep 3: Extract SWES 1 scenario params (from project docs).\nStep 4: LBO market benchmarks (S&P LCD, Fund Finance Association).\n\nOutputs: sector_elasticities.json, scenario_params.json, market_reference.json')

doc.add_heading('WS3: Agent Balance Sheet Sizing (Analyst 2, Days 4-5)', level=3)
doc.add_paragraph('Goal: Calibrate balance sheet proportions per agent type.\n\nSources: BoE Bankstats, BoE FSR, TPR annual reports, SWES 2 sizing data, Fund Finance Association surveys.\n\nOutput: agent_calibration.json')

doc.add_page_break()

# ═══════ 9. BUILD TIMELINE ═══════
doc.add_heading('9. Build Timeline', level=1)
add_table(['Day', 'Developer', 'Analyst', 'Gate'], [
    ['1', 'Repo setup, dataclasses, data_loader.py', 'WS1: identify companies, start CH extractions', ''],
    ['2', 'scenario_engine.py: stress_company(), apply_scenario', 'WS1 cont; WS2: ONS data, SWES 1 params', ''],
    ['3', 'agents.py: factory functions for Bank, AAM, PF', 'WS1 complete (40+ cos); WS2 cont', ''],
    ['4', 'feedback_loop.py: reactions, 2nd round, reputation', 'WS2 complete; WS3: agent sizing', 'Engine testable'],
    ['5', 'Streamlit: sidebar, Tab 1 (Portfolio), Tab 2 (Direct)', 'WS3 complete; data QA', ''],
    ['6-7', 'Tab 3 (System Dynamics + toggle), Tab 4 (Real Econ)', 'Demo dataset variants', ''],
    ['8', 'Integration testing, bug fixes, performance', '', 'v0.1 ready'],
    ['9-10', 'INTERNAL DEMO. Collect feedback.', '', 'INTERNAL DEMO'],
    ['11-14', 'Polish, drill-downs, amplification calibration', 'Client-specific portfolio prep', ''],
    ['15-18', 'Final polish, demo rehearsal, narrative scripts', '', ''],
    ['19-20', 'CLIENT DEMO', '', 'CLIENT DEMO'],
])

doc.add_page_break()

# ═══════ 10. RISK REGISTER ═══════
doc.add_heading('10. Risk Register', level=1)
add_table(['Risk', 'L', 'I', 'Mitigation'], [
    ['Feedback loop = black box', 'H', 'H', 'Calibrate to SWES 1 published findings + Van den End (2009). Prepare equation reference as leave-behind.'],
    ['Portfolio looks synthetic', 'M', 'H', 'Use real UK PE companies from Companies House. Exposure amounts synthetic, borrower universe authentic.'],
    ['Amplification ratio wrong', 'M', 'M', 'Target 1.5-3.0x (Van den End = 3.0x). Calibrate channel intensities iteratively.'],
    ['Companies House data insufficient', 'M', 'L', 'Fallback: synthetic cos from Preqin/BVCA averages. Label "illustrative portfolio."'],
    ['Deep methodology questions', 'M', 'M', '5-page methodology appendix as leave-behind (not shown in demo).'],
    ['Streamlit perf with live sliders', 'L', 'M', 'Pre-compute scenario grid; st.cache_data(). Target <2s response.'],
])

doc.add_page_break()

# ═══════ APPENDIX A ═══════
doc.add_heading('Appendix A: SWES 1 Scenario Parameters', level=1)
add_table(['Variable', 'SWES 1 Value', 'Comparison'], [
    ['10Y Gilt Yield', '+115 bps / 10 days', '~90% of Sep 2022 LDI episode'],
    ['Sterling IG Spread', '+130 bps', '~Equal to Mar 2020 dash for cash'],
    ['10Y UST Yield', 'Comparable to largest since 2001', 'Global shock'],
    ['Equity', 'Less than historical max', 'Not dominant shock'],
    ['Narrative: HF default', 'Mid-sized RV fund', 'Counterparty credit risk'],
    ['Narrative: Downgrade', 'Single notch, UK + others', 'Broad credit risk'],
    ['Narrative: SWF selling', 'Reduce adv. economy debt', 'Gilt selling pressure'],
])

doc.add_page_break()

# ═══════ APPENDIX B ═══════
doc.add_heading('Appendix B: Van den End Equation Mapping', level=1)
add_table(['VdE Eq', 'Original', 'Our Adaptation'], [
    ['(1)', 'B\u2080 = \u03a3 I_non_cal,i', 'Agent-specific buffer defs'],
    ['(2)', 'E\u2081 = \u03a3 I\u1d62 \u00d7 w_sim\u2081,\u1d62', 'Weights from company stress (2b-2f) not MC'],
    ['(3)', 'B\u2081 = B\u2080 \u2212 E\u2081', 'Unchanged'],
    ['(4)', 'RI\u1d62 = (B\u2080\u2212B\u2081)(I\u1d62/\u03a3I)', 'Unchanged; agent-specific instruments'],
    ['(5)', 'B\u2082 = B\u2081 + \u03a3 RI\u1d62(1\u2212w\u2081/100)', 'Unchanged'],
    ['(6)', 'w\u2082 = f(w\u2081, \u03a3q, similarity, s)', 'Decomposed into cross-agent channels'],
    ['(7)', 'w*\u2082 = w\u2082 \u00d7 \u221as', 'Unchanged; all reacting agents'],
    ['(8)', 'E\u2082 = \u03a3(I+RI)(w\u2082\u2212w\u2081)', 'Includes cross-agent effects'],
    ['(9)', 'B\u2083 = B\u2082 \u2212 E\u2082', 'Unchanged'],
])

doc.add_paragraph('\nKey VdE result: Credit crisis amplification ratio ~3.0x (39% total loss vs 13% first round). Reputation risk = >50% of second round. No correlation between shortfall probability and bank size in second round = systemic risk indicator.')

doc.add_page_break()

# ═══════ APPENDIX C: DEMO SCRIPTS ═══════
doc.add_heading('Appendix C: Demo Narrative Scripts', level=1)

doc.add_heading('Script A: Bank Client Framing', level=2)
doc.add_paragraph('Opening: "The BoE\'s second SWES is focused on private markets \u2014 your exposure to PE funds, leveraged lending, subscription lines, NAV facilities. Unlike a conventional stress test, SWES asks how your actions interact with other market participants to amplify the shock. We\'ve built a tool that models exactly this."')
doc.add_paragraph('Act 3 Turn: "But the BoE isn\'t asking about direct stress. They\'re asking about system-wide dynamics. Let me flip this switch. [Toggle ON.] Watch what happens. Your bank rationally tightens lending \u2014 exactly what SWES 1 showed banks do. But when every bank does this simultaneously, PE companies needing refinancing face a cliff. Defaults spike. Leveraged loan market freezes. CLO OC tests fail. Your losses just went from \u00a3X to \u00a32.5X. The amplification ratio is 2.3x \u2014 consistent with DNB\'s published research showing ratios up to 3x."')
doc.add_paragraph('Close: "We calibrate this to your actual portfolio, design your SWES 2 submission methodology, and bring cross-market insight on how AAMs and investors will behave."')

doc.add_heading('Script B: AAM Client Framing', level=2)
doc.add_paragraph('Opening: "SWES 2 asks you to model how a prolonged downturn affects your UK portfolio companies \u2014 and critically, how banks, investors, and other managers collectively amplify that stress. This is new territory for most AAMs."')
doc.add_paragraph('Act 3 Turn: "SWES 1 proved that banks are significantly more conservative than fund managers expect. Over a third of managers expecting financing were refused by every bank. When we model this tightening, your portfolio companies needing refinancing face a cliff. Covenant breaches jump from Y to Z. And your LPs are simultaneously facing their own pressures, so new commitments dry up."')
doc.add_paragraph('Close: "We work with banks on their side of this equation, giving us unique insight into how they\'ll behave. We help you build a SWES 2 submission that accounts for these dynamics."')

# ═══════ SAVE ═══════
out = "/mnt/user-data/outputs/SWES_Private_Markets_Stress_Lens_FRD.docx"
doc.save(out)
print(f"Saved to {out}")
print(f"{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
